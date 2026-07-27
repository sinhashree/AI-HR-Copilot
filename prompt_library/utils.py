from io import BytesIO
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet


def create_word(text):
    """
    Create a Word document and return it as bytes.
    """
    document = Document()

    document.add_heading("AI Generated Job Description", level=1)
    document.add_paragraph(text)

    file = BytesIO()
    document.save(file)
    file.seek(0)

    return file


def create_pdf(text):
    """
    Create a PDF document and return it as bytes.
    """
    file = BytesIO()

    doc = SimpleDocTemplate(file)

    styles = getSampleStyleSheet()

    story = []

    for line in text.split("\n"):
        story.append(Paragraph(line, styles["BodyText"]))

    doc.build(story)

    file.seek(0)

    return file